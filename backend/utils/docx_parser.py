import logging
from typing import List, Optional
import re

logger = logging.getLogger(__name__)

class DOCXParser:
    """
    DOCX Parser with paragraph-aware chunking
    
    Features:
    - Text extraction from DOCX
    - Paragraph-aware chunking
    - Metadata extraction
    """
    
    def __init__(self, chunk_size: int = 500, overlap: int = 50):
        """
        Initialize DOCX Parser
        
        Args:
            chunk_size: Target size of each chunk in tokens (approximate)
            overlap: Number of tokens to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.overlap = overlap
    
    def extract_text(self, file_path: str) -> Optional[str]:
        """
        Extract raw text from DOCX
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Extracted text or None if failed
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            paragraphs = []
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            full_text = "\n".join(paragraphs)
            logger.info(f"Extracted {len(full_text)} characters from {len(paragraphs)} paragraphs")
            return full_text
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX {file_path}: {e}")
            return None
    
    def extract_text_with_chunking(self, file_path: str) -> List[str]:
        """
        Extract text from DOCX and split into chunks
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            List of text chunks
        """
        text = self.extract_text(file_path)
        if not text:
            return []
        
        return self.chunk_text(text)
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks, respecting paragraph boundaries
        
        Args:
            text: Input text
            
        Returns:
            List of text chunks
        """
        # Split by paragraphs first
        paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
        
        if not paragraphs:
            return []
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for para in paragraphs:
            para_length = len(para.split())
            
            # If paragraph alone exceeds chunk size, split it further
            if para_length > self.chunk_size:
                # Save current chunk if exists
                if current_chunk:
                    chunks.append("\n".join(current_chunk))
                    current_chunk = []
                    current_length = 0
                
                # Split large paragraph into sentences
                sentences = self._split_into_sentences(para)
                sub_chunk = []
                sub_length = 0
                
                for sentence in sentences:
                    s_len = len(sentence.split())
                    if sub_length + s_len > self.chunk_size and sub_chunk:
                        chunks.append(" ".join(sub_chunk))
                        
                        # Overlap
                        overlap_text = []
                        overlap_len = 0
                        for s in reversed(sub_chunk):
                            s_l = len(s.split())
                            if overlap_len + s_l <= self.overlap:
                                overlap_text.insert(0, s)
                                overlap_len += s_l
                            else:
                                break
                        
                        sub_chunk = overlap_text
                        sub_length = overlap_len
                    
                    sub_chunk.append(sentence)
                    sub_length += s_len
                
                if sub_chunk:
                    chunks.append(" ".join(sub_chunk))
                
                continue
            
            # If adding this paragraph exceeds chunk size, save current chunk
            if current_length + para_length > self.chunk_size and current_chunk:
                chunks.append("\n".join(current_chunk))
                
                # Create overlap
                overlap_paras = []
                overlap_length = 0
                for p in reversed(current_chunk):
                    p_len = len(p.split())
                    if overlap_length + p_len <= self.overlap:
                        overlap_paras.insert(0, p)
                        overlap_length += p_len
                    else:
                        break
                
                current_chunk = overlap_paras
                current_length = overlap_length
            
            current_chunk.append(para)
            current_length += para_length
        
        # Add final chunk
        if current_chunk:
            chunks.append("\n".join(current_chunk))
        
        logger.info(f"Created {len(chunks)} chunks from text")
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences"""
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def extract_metadata(self, file_path: str) -> dict:
        """
        Extract metadata from DOCX
        
        Args:
            file_path: Path to DOCX file
            
        Returns:
            Dictionary of metadata
        """
        try:
            from docx import Document
            
            doc = Document(file_path)
            metadata = {
                'file_path': file_path,
                'num_paragraphs': len(doc.paragraphs)
            }
            
            # Extract core properties if available
            if hasattr(doc, 'core_properties'):
                props = doc.core_properties
                metadata.update({
                    'title': props.title or '',
                    'author': props.author or '',
                    'subject': props.subject or '',
                    'created': str(props.created) if props.created else ''
                })
            
            return metadata
            
        except Exception as e:
            logger.error(f"Failed to extract metadata from DOCX: {e}")
            return {'file_path': file_path}

def parse_docx(file_storage) -> str:
    """
    Wrapper function to parse DOCX from Flask FileStorage object
    """
    try:
        from docx import Document
        
        # Document accepts file path or file-like object
        doc = Document(file_storage)
        
        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)
                
        return "\n".join(paragraphs)
        
    except Exception as e:
        logger.error(f"Error in parse_docx wrapper: {e}")
        return ""
