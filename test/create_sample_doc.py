
from docx import Document

def create_sample_docx():
    doc = Document()
    doc.add_heading('Introduction to Artificial Intelligence', 0)
    
    doc.add_paragraph(
        "Artificial Intelligence (AI) is intelligence demonstrated by machines, "
        "as opposed to the natural intelligence displayed by humans and animals."
    )
    
    doc.add_heading('Machine Learning', level=1)
    doc.add_paragraph(
        "Machine learning (ML) is a field of inquiry devoted to understanding and "
        "building methods that 'learn', that is, methods that leverage data to "
        "improve performance on some set of tasks."
    )
    
    doc.add_heading('Neural Networks', level=1)
    doc.add_paragraph(
        "Artificial neural networks (ANNs) are computing systems inspired by the "
        "biological neural networks that constitute animal brains."
    )
    
    doc.save('test/sample_test.docx')
    print("Created test/sample_test.docx")

if __name__ == "__main__":
    create_sample_docx()
